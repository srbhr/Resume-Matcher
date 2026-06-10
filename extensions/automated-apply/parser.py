import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_text_from_docx(file_path):
    """
    Extracts all text from a docx file.
    Returns it as a single string.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading docx: {e}")
        return ""

def add_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # Add bottom border logic here if possible, or just a line
    # docx doesn't easily do bottom borders on paragraphs natively, 
    # so we'll just keep it bold and centered.
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    
    # Fake a border line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("_" * 70)
    run_line.font.size = Pt(8)
    run_line.font.name = 'Times New Roman'
    p_line.paragraph_format.space_after = Pt(4)
    p_line.paragraph_format.space_before = Pt(0)

def create_tailored_docx(resume_json, output_path):
    """
    Creates a new docx file with Ivy League formatting based on strict JSON.
    """
    if not resume_json:
        print("No JSON provided to parser.")
        return None

    try:
        doc = docx.Document()
        
        # Set even narrower margins (0.4 inches) to save space
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)
            
        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        # 1. Contact Header
        contact = resume_json.get('contact', {})
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(contact.get('name', 'NAME').upper())
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_p.paragraph_format.space_after = Pt(2)
        
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links = []
        for k in ['email', 'phone', 'linkedin', 'github', 'leetcode']:
            if contact.get(k):
                links.append(contact[k])
        
        info_run = info_p.add_run(" | ".join(links))
        info_run.font.size = Pt(10)
        info_p.paragraph_format.space_after = Pt(8)

        # 2. Summary
        if resume_json.get('summary'):
            add_header(doc, "Summary")
            sum_p = doc.add_paragraph(resume_json['summary'])
            sum_p.paragraph_format.space_after = Pt(2)

        # 3. Education
        if resume_json.get('education'):
            add_header(doc, "Education")
            for edu in resume_json['education']:
                edu_p = doc.add_paragraph()
                edu_run1 = edu_p.add_run(f"{edu.get('university', '')} ")
                edu_run1.bold = True
                
                # We can simulate right alignment with a tab, but for simplicity:
                edu_p.add_run(f"- {edu.get('degree', '')} (CGPA: {edu.get('cgpa', '')}) | {edu.get('graduation_year', '')}")
                edu_p.paragraph_format.space_after = Pt(2)

        # 4. Projects
        if resume_json.get('projects'):
            add_header(doc, "Projects")
            for proj in resume_json['projects']:
                p_title = doc.add_paragraph()
                r_title = p_title.add_run(proj.get('name', ''))
                r_title.bold = True
                
                if proj.get('technologies'):
                    r_tech = p_title.add_run(f" | {proj['technologies']}")
                    r_tech.italic = True
                    
                p_title.paragraph_format.space_after = Pt(2)
                
                for bullet in proj.get('bullets', []):
                    b_p = doc.add_paragraph(bullet, style='List Bullet')
                    b_p.paragraph_format.space_after = Pt(0)
                    b_p.paragraph_format.left_indent = Inches(0.25)
                
                # Small gap after project
                gap = doc.add_paragraph()
                gap.paragraph_format.space_after = Pt(0)

        # 5. Skills
        if resume_json.get('skills'):
            add_header(doc, "Skills")
            skills = resume_json['skills']
            for k, v in skills.items():
                if v:
                    s_p = doc.add_paragraph()
                    s_name = s_p.add_run(f"{k.capitalize()}: ")
                    s_name.bold = True
                    s_p.add_run(v)
                    s_p.paragraph_format.space_after = Pt(2)

        doc.save(output_path)
        print(f"Successfully saved Ivy League formatted resume to {output_path}")
        return output_path
    except Exception as e:
        print(f"Error creating tailored docx: {e}")
        return None
